"""Exporters for the merged transcript: TXT, SRT, JSON, DOCX."""
from __future__ import annotations

import json
from pathlib import Path

from app.merge import Chunk


def _fmt_hms(seconds: float) -> str:
    seconds = max(0.0, seconds)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _fmt_srt_time(seconds: float) -> str:
    seconds = max(0.0, seconds)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int(round((seconds - int(seconds)) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def to_txt(chunks: list[Chunk], path: str) -> None:
    lines = [f"[{_fmt_hms(c.start)} - {_fmt_hms(c.end)}] {c.speaker}: {c.text}" for c in chunks]
    Path(path).write_text("\n\n".join(lines), encoding="utf-8")


def to_srt(chunks: list[Chunk], path: str) -> None:
    blocks = []
    for i, c in enumerate(chunks, start=1):
        blocks.append(
            f"{i}\n{_fmt_srt_time(c.start)} --> {_fmt_srt_time(c.end)}\n{c.speaker}: {c.text}\n"
        )
    Path(path).write_text("\n".join(blocks), encoding="utf-8")


def to_json(chunks: list[Chunk], path: str) -> None:
    data = [
        {"start": round(c.start, 2), "end": round(c.end, 2), "speaker": c.speaker, "text": c.text}
        for c in chunks
    ]
    Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def to_docx(chunks: list[Chunk], path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    for c in chunks:
        p = doc.add_paragraph()
        run = p.add_run(f"[{_fmt_hms(c.start)} - {_fmt_hms(c.end)}] {c.speaker}: ")
        run.bold = True
        p.add_run(c.text)
    doc.save(path)
